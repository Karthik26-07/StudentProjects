from django.shortcuts import render
from PersonalityApp.models import *
from django.db.models import Q
import sweetify
from django.http import HttpResponse
import json
from django.core.files.storage import FileSystemStorage
import random
import base64
import codecs
from PIL import Image
import io
from docxtpl import DocxTemplate
import os
from docx import Document
# import aspose.words as aw
# import docx2txt

# Create your views here.


def index(request):
    return render(request,'index.html')

def user(request):
    return render(request,'user/index.html')


def registration(request):
    return render(request,'user/registration.html')


def saveUser(request):
    if request.method == 'POST':
        farmername = request.POST['uname']
        contactNo = request.POST['contactNo']
        emailId = request.POST['emailId']
        address = request.POST['address']
        username = request.POST['username']
        password = request.POST['password']
        
        CVDocumnet = request.FILES['docfile'].read()
       
        
        cv=bytearray(CVDocumnet)
        # print(CVDocumnet.decode("latin1").isascii())
        # stream= cv.decode("latin1")
        # stream=int.from_bytes(cv, byteorder='big', signed=False)
        filename = "Document-"+str(random.randint(1000000000, 9999999999))+".doc"
        # res=""
        a = base64.b64encode(cv)
        with open('Documents/'+filename, 'wb') as f:
          f.write(base64.b64decode(a))
       
        # s=base64.b64encode('welcome')
        # print(base64.b64decode(s))
       
        # document = Document()
        # # # document.add_heading('A simple text', level=1)
        
        # stream= cv.decode("latin1")
        # print(type(stream))
        # # print(stream)
        
        # document.add_paragraph(CVDocumnet.decode("latin1"))
       
        
        # # document = Document(io.BytesIO(CVDocumnet))
        # document.save('Documents/'+filename)
        
       
        
        user = User.objects.filter(
            Q(email=emailId) | Q(contact=contactNo) | Q(user_name=username)
        ).first()

        has_error = False
        error = ''

        if user != None and user.user_name == username:
            has_error = True
            error = 'Duplicate user name'

        if user != None and user.email == emailId:
            has_error = True
            error = 'Duplicate email'

        if user != None and user.contact == contactNo:
            has_error = True
            error = 'Duplicate contact number'

        if has_error:
            return render(request, "user/registration.html", {'error': error})

        user = User(name=farmername, contact=contactNo, email=emailId,
                    address=address, user_name=username, password=password)
        user.save()

        return render(request, "user/registration.html", {'success': 'User Added Successfully'})
    else:
        return render(request, 'user/registration.html')


def userlogin(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = User.objects.values_list('password', 'id', 'name').\
            filter(user_name=request.POST['username'])

        user = User.objects.filter(
            user_name=username, password=password).first()

        if user == None:
            return render(request, 'user/index.html', {'error': 'Invalid login credentials'})

        request.session['userid'] = user.id
        request.session['userName'] = user.name

        return render(request, 'user/userHome.html')

    else:
        return render(request, 'user/index.html')


def adminlogin(request):
    return render(request,'admin/login.html')

def adminloginview(request):
    if request.method == "POST":
        adminusername=request.POST["username"]
        adminpass=request.POST["password"]

        if adminusername == 'admin' and adminpass == 'admin':
            return render(request,'admin/home.html')
        else:
            return render(request, 'admin/login.html', {'error': 'Invalid login credentials'})

def Apptitude(request):
    return render(request,'admin/AptitudeQuestions.html')

def Add_Apptitude(request):
    if request.method=="POST":
         Question = request.POST["Question"]
         option1 = request.POST["option1"]
         option2 = request.POST["option2"]
         option3 = request.POST["option3"]
         option4 = request.POST["option4"]
         Answer = request.POST["correctAns"]
         Type=request.POST["type"]
         data=apptitude_question(questionType=Type,questions=Question,
         option1=option1,option2=option2,option3=option3,option4=option4,answer=Answer)
         data.save()
         sweetify.success(request, 'Added', text='Successfully Added ', persistent='OK')
    return render(request,'admin/AptitudeQuestions.html')
      
def Pesonality(request):
    return render(request,'admin/PersonalityQuestion.html')

def Add_Pesonality(request):
    if request.method=="POST":
         question = request.POST["Question"]
         openness_to_experience = request.POST["Openness"]
         conscientiousness = request.POST["Conscientiousness"]
         extraversion = request.POST["Extraversion"]
         agreeableness = request.POST["Agreebleness"]
         neuroticism = request.POST["Neuroticism"]
       
         data=personality_question(question=question,openness_to_experience=openness_to_experience,
         conscientiousness=conscientiousness,extraversion=extraversion,
         agreeableness=agreeableness,neuroticism=neuroticism)
         data.save()
         sweetify.success(request, 'Added', text='Successfully Added ', persistent='OK')
    return render(request,'admin/PersonalityQuestion.html')


def displayAptitude(request):
    aptitude=apptitude_question.objects.all().values()
    return render(request,'admin/DisplayAptitudeQuestion.html',{'Aptitude':aptitude}) 

def DisplayPersonality(request):
    personality=personality_question.objects.all().values()

    return render(request,'admin/DisplayPersonalityQuestion.html',{'Personality':personality}) 

def personalityTest(request):
      aptitude=apptitude_question.objects.all().values()
      personality=personality_question.objects.all().values()
      return render(request, 'user/personalityTest.html',{'Personality':personality,'Aptitude':aptitude})

def aptitudeTest(request):
      aptitude=apptitude_question.objects.all().values()
      personality=personality_question.objects.all().values()
      return render(request, 'user/aptitudeTest.html',{'Personality':personality,'Aptitude':aptitude})


def Add_Job(request):
    
    return render(request,'admin/jobDetails.html') 

def Save_Job(request):
      if request.method=="POST":
        
         Designation = request.POST["designation"]
         Place = request.POST["place"]
         Salary = request.POST["salary"]
         jobid = request.POST["jobid"]
         data=job_details(jobid=jobid,designation=Designation,salary=Salary,place=Place)
         data.save()
         sweetify.success(request, 'Added', text='Successfully Added ', persistent='OK')
   
      return render(request,'admin/jobDetails.html') 

def jobRequirement(request):
    
    data=job_details.objects.all().values()
    return render(request,'admin/jobRequirement.html',{'jobdata':data}) 

def jobdetails(request):
    if request.method=="GET":
        data=[] 
        ids=request.GET["id"]
        jobdata=job_details.objects.all().filter(id=ids).values()
        data.extend(jobdata)
        print(data)
 
    return HttpResponse(json.dumps(data), content_type="application/json")
    
def Add_Job_Requirement(request):
    if request.method =="POST":
        Experience=request.POST['Experience']
        Qualifiction=request.POST['Qualifiction']
        keySkills=request.POST['keySkills']
        Jobid=request.POST['Jobid']
        data=job_details.objects.all().values()
        job_req=job_requirement(experience=Experience,job_id=Jobid,qualification=Qualifiction,skills=keySkills)
        job_req.save()
        sweetify.success(request, 'Added', text='Successfully Added ', persistent='OK')
    
    return render(request,'admin/jobRequirement.html',{'jobdata':data})     


def displayJobDetail(request):
    data=job_details.objects.all().values()
    return render(request,'admin/DisplayJobDetails.html',{'jobdetails':data}) 

def UploadPreferedCV(request):
    
    data=job_details.objects.all().values()
    return render(request,'admin/UploadPreferedCV.html',{'jobdata':data}) 

def EditPersonality(request):
        return render(request,'admin/UploadPreferedCV.html') 

def EditViewDetails(request):
    if request.method=="GET":
      command=request.GET['command']
      ids=request.GET['id']
      data=[]
      if command =='personality':
         Personality=personality_question.objects.all().filter(id=ids).values()
         data.extend(Personality)
      elif command =='apptitude': 
         aptitude=apptitude_question.objects.all().filter(id=ids).values()
         data.extend(aptitude) 
      elif command =='JobDetails':
         job=job_details.objects.all().filter(id=ids).values()
         data.extend(job) 

    
    return HttpResponse(json.dumps(data), content_type="application/json") 


def UpdatePersonality(request):
    if request.method=="POST":
    
         question = request.POST["Question"]
         openness_to_experience = request.POST["Openness"]
         conscientiousness = request.POST["Conscientiousness"]
         extraversion = request.POST["Extraversion"]
         agreeableness = request.POST["Agreebleness"]
         neuroticism = request.POST["Neuroticism"]
         ID=request.POST['id']

         personality_question.objects.filter(id=ID).update(question=question,openness_to_experience=openness_to_experience,
         conscientiousness=conscientiousness,extraversion=extraversion,
         agreeableness=agreeableness,neuroticism=neuroticism)
          
         personality=personality_question.objects.all().values()

         sweetify.success(request, 'Success', text='Successfully Updated ', persistent='OK')

    return render(request,'admin/DisplayPersonalityQuestion.html',{'Personality':personality}) 
         

def DeletePersonality(request): 
    if request.method=="GET":  
        id=request.GET['id']
        personality_question.objects.filter(id=id).delete()
    
    
    return render(request,'admin/DisplayPersonalityQuestion.html') 
         
def UpdateAptitude(request):
    if request.method=="POST":
         Question = request.POST["Question"]
         option1 = request.POST["option1"]
         option2 = request.POST["option2"]
         option3 = request.POST["option3"]
         option4 = request.POST["option4"]
         Answer = request.POST["correctAns"]
         Type=request.POST["type"]
         ID=request.POST['id']
         apptitude_question.objects.filter(id=ID).update(questionType=Type,questions=Question,
         option1=option1,option2=option2,option3=option3,option4=option4,answer=Answer)
         
         aptitude=apptitude_question.objects.all().values()
         sweetify.success(request, 'Success', text='Successfully Updated ', persistent='OK')

    return render(request,'admin/DisplayAptitudeQuestion.html',{'Aptitude':aptitude}) 

def DeleteAptitude(request):
     if request.method=="GET":  
        id=request.GET['id']
        apptitude_question.objects.filter(id=id).delete()
    
    
     return render(request,'admin/DisplayAptitudeQuestion.html') 
         
def UpdateJobDetail(request):  
     if request.method=="POST":
        
         Designation = request.POST["designation"]
         Place = request.POST["place"]
         Salary = request.POST["salary"]
         jobid = request.POST["jobid"]  
         ID=request.POST['id']   
         
         job_details.objects.filter(id=ID).update(jobid=jobid,designation=Designation,salary=Salary,place=Place)  
         data=job_details.objects.all().values()
         sweetify.success(request, 'Success', text='Successfully Updated ', persistent='OK')
    
     return render(request,'admin/DisplayJobDetails.html',{'jobdetails':data}) 

def DeleteJobDetail(request):
     if request.method=="GET":  
        id=request.GET['id']
        job_details.objects.filter(id=id).delete()
    
    
     return render(request,'admin/DisplayJobDetails.html') 